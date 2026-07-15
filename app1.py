import streamlit as st
from google import genai
from config import API_KEY

from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
import io

st.set_page_config(
    page_title="AI Resume Builder",
    page_icon="📄",
    layout="wide"
)

if not API_KEY:
    st.error("🔑 Gemini API Key is missing!")
    st.info("""
    To run this app on Streamlit Cloud, you need to set your Gemini API Key in the Streamlit Secrets:
    
    1. Go to your **Streamlit App Dashboard**.
    2. Click on the app's settings (**Settings** -> **Secrets**).
    3. Add the following:
       ```toml
       GOOGLE_GEMINI_API_KEY = "your_actual_api_key_here"
       ```
    4. Save the secrets and reboot the app.
    
    *If running locally, please add `GOOGLE_GEMINI_API_KEY=your_key` to a `.env` file in the root directory.*
    """)
    st.stop()

client = genai.Client(api_key=API_KEY)

st.title("📄 AI Resume Builder")

st.sidebar.title("Settings")

resume_type = st.sidebar.selectbox(
    "Resume Type",
    [
        "Fresher",
        "Experienced",
        "ATS Friendly"
    ]
)

photo = st.file_uploader(
    "Upload Profile Photo (Optional)",
    type=["jpg", "jpeg", "png"]
)

with st.form("resume"):

    name = st.text_input("Full Name")
    email = st.text_input("Email")
    phone = st.text_input("Phone")
    address = st.text_area("Address")

    objective = st.text_area("Career Objective")

    education = st.text_area("Education")

    skills = st.text_area("Skills")

    projects = st.text_area("Projects")

    experience = st.text_area("Experience")

    certifications = st.text_area("Certifications")

    achievements = st.text_area("Achievements")

    languages = st.text_input("Languages")

    generate = st.form_submit_button("Generate Resume")

if generate:

    prompt = f"""
Create a professional {resume_type} resume.

Name:
{name}

Email:
{email}

Phone:
{phone}

Address:
{address}

Career Objective:
{objective}

Education:
{education}

Skills:
{skills}

Projects:
{projects}

Experience:
{experience}

Certifications:
{certifications}

Achievements:
{achievements}

Languages:
{languages}

Use proper headings.
Use bullet points.
Return only the resume.
"""

    with st.spinner("Generating Resume..."):

        response = client.models.generate_content(
            model="gemini-3.5-flash",
            contents=prompt
        )

    resume = response.text

    st.success("Resume Generated Successfully")

    if photo:
        st.image(photo, width=150)

    st.markdown(resume)

    # Download TXT

    st.download_button(
        "Download TXT",
        resume,
        "resume.txt",
        "text/plain"
    )

    # Create Word File in memory

    doc = Document()
    doc.add_heading("Resume", level=1)
    doc.add_paragraph(resume)

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)

    st.download_button(
        "Download Word",
        doc_io,
        "resume.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    # Create PDF in memory

    styles = getSampleStyleSheet()
    pdf_io = io.BytesIO()
    pdf = SimpleDocTemplate(pdf_io)
    story = []

    for line in resume.split("\n"):
        story.append(Paragraph(line, styles["BodyText"]))

    pdf.build(story)
    pdf_io.seek(0)

    st.download_button(
        "Download PDF",
        pdf_io,
        "resume.pdf",
        "application/pdf"
    )